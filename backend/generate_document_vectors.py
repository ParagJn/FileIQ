"""
vector_db_manager.py

A reusable, production-grade class for chunking and embedding text from JSON, PDF, and DOCX files, and managing a vector database using FAISS and SentenceTransformers. 

This module provides a robust, memory-efficient, and extensible solution for semantic search and vector-based retrieval over various file types. It supports advanced features such as batch embedding with caching, optimized chunking, HTML cleaning, and integration with large language models (LLMs) for enhanced Q&A.

Features:
- Efficiently loads and processes JSON, PDF, and DOCX files with error handling and progress logging using tqdm progress bars.
- Cleans and chunks text for optimal embedding and retrieval.
- Uses SentenceTransformers for high-quality embeddings with batch processing and LRU cache.
- Manages a FAISS-based vector database with support for building, loading, rebuilding, and optimizing indices.
- Provides semantic search with optional score filtering and parent mapping for context.
- Includes interactive and automated demo/test functions with Claude Sonnet LLM integration for enhanced document Q&A.
- Visual progress tracking with tqdm for all long-running operations (PDF processing, embedding generation, batch operations).

Author: Parag Jain
Date: 2025-06-19
"""
import os
import json
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer
import pickle
import hashlib
import threading
from typing import Dict, List, Tuple, Optional, Any
import logging
from functools import lru_cache
from tqdm import tqdm

class VectorDBManager:
    """
    VectorDBManager
    ----------------
    Manages chunking and embedding of text from JSON, PDF, and DOCX files, and building/loading/rebuilding a vector database for semantic search.

    Args:
        folder_path (str): Path to the folder containing files. This is the entry point for the vectorization process.
        vector_data_dir (str, optional): Directory to store vector DB files. Defaults to a 'vector-data' subfolder.
        model_name (str, optional): Name of the SentenceTransformer model to use for embeddings.
        batch_size (int, optional): Batch size for embedding generation. Defaults to 32.
        cache_size (int, optional): Maximum number of embeddings to cache in memory. Defaults to 1000.

    Main Methods:
        load_file(filename): Load and extract text from JSON, PDF, or DOCX file.
        remove_html(obj): Recursively remove HTML tags from all string values in a JSON-like object.
        text_to_chunks(text, ...): Split plain text into overlapping chunks for embedding.
        json_to_chunks(obj, ...): Recursively split JSON data into text chunks for embedding.
        encode_batch(texts): Batch-encode texts into embeddings with caching.
        build_vector_db(filename, data=None): Build a new vector database from a file or data.
        load_vector_db(filename): Load an existing vector database and validate its integrity.
        search_vectors(filename, query, ...): Search the vector database for relevant chunks given a query.
        delete_vectors(filename, chunk_indices): Delete specific vectors from the database.
        get_database_stats(filename): Get statistics about the vector database.
        optimize_index(filename, ...): Optimize the FAISS index for better search performance.
        rebuild_vector_db(filename): Delete and rebuild the vector database for a file.

    Usage Example:
        db_manager = VectorDBManager(folder_path="/path/to/files")
        db_manager.build_vector_db("mydoc.pdf")
        results = db_manager.search_vectors("mydoc.pdf", "What is the summary?")
    """
    def __init__(self, folder_path, vector_data_dir=None, model_name='all-MiniLM-L6-v2', batch_size=32, cache_size=1000):
        self.folder_path = folder_path
        self.vector_data_dir = vector_data_dir or os.path.join(folder_path, 'vector-data')
        os.makedirs(self.vector_data_dir, exist_ok=True)
        self.model_name = model_name
        self._model = None
        self._model_lock = threading.Lock()
        self.batch_size = batch_size
        self.cache_size = cache_size
        self._embedding_cache = {}
        self._setup_logging()

    def load_file(self, filename):
        """
        Loads and extracts text from a file with optimized memory usage and error handling.
        Args:
            filename (str): Name of the file.
        Returns:
            dict or list: Parsed JSON data or converted structure.
        Raises:
            RuntimeError: If the file cannot be loaded or parsed.
        """
        ext = os.path.splitext(filename)[1].lower()
        file_path = os.path.join(self.folder_path, filename)
        
        # Check file exists and is readable
        if not os.path.exists(file_path):
            raise RuntimeError(f"File not found: {filename}")
        
        if not os.access(file_path, os.R_OK):
            raise RuntimeError(f"File not readable: {filename}")
        
        # Check file size and warn for large files
        file_size = os.path.getsize(file_path)
        if file_size > 100 * 1024 * 1024:  # 100MB
            self.logger.warning(f"Processing large file ({file_size / (1024*1024):.1f} MB): {filename}")
        
        try:
            if ext == '.json':
                with open(file_path, 'r', encoding='utf-8', buffering=8192) as f:
                    return json.load(f)
                    
            elif ext == '.pdf':
                try:
                    import PyPDF2
                except ImportError:
                    raise RuntimeError("PyPDF2 is required to process PDF files.")
                    
                text_sections = []
                with open(file_path, 'rb') as f:
                    try:
                        reader = PyPDF2.PdfReader(f)
                        total_pages = len(reader.pages)
                        
                        # Use tqdm progress bar for PDF processing
                        with tqdm(total=total_pages, desc=f"Processing {filename}", unit="pages") as pbar:
                            for i, page in enumerate(reader.pages):
                                page_text = page.extract_text() or ""
                                if page_text.strip():
                                    # Split by double newlines or paragraphs
                                    sections = [s.strip() for s in page_text.split('\n\n') if s.strip()]
                                    text_sections.extend(sections)
                                
                                pbar.update(1)
                                
                    except Exception as e:
                        raise RuntimeError(f"Failed to parse PDF: {e}")
                        
                return {"sections": text_sections}
                
            elif ext == '.docx':
                try:
                    import docx
                except ImportError:
                    raise RuntimeError("python-docx is required to process DOCX files.")
                    
                try:
                    doc = docx.Document(file_path)
                    chunks = []
                    current_chunk = []
                    
                    def flush_chunk():
                        if current_chunk:
                            chunk_text = "\n".join(current_chunk).strip()
                            if chunk_text:
                                chunks.append(chunk_text)
                            current_chunk.clear()
                    
                    # Process paragraphs
                    for para in doc.paragraphs:
                        if not para.text.strip():
                            continue
                            
                        style = para.style.name if hasattr(para.style, 'name') else str(para.style)
                        
                        if style.startswith('Heading'):
                            flush_chunk()
                            current_chunk.append(f"[HEADER] {para.text.strip()}")
                        else:
                            current_chunk.append(para.text.strip())
                            
                        # Flush chunk if it gets too large
                        if len("\n".join(current_chunk)) > 2000:
                            flush_chunk()
                    
                    flush_chunk()
                    
                    # Process tables separately
                    for table in doc.tables:
                        table_text = []
                        for row in table.rows:
                            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                            if row_text:
                                table_text.append("\t".join(row_text))
                        
                        if table_text:
                            chunks.append("[TABLE]\n" + "\n".join(table_text))
                    
                    return {"chunks": chunks}
                    
                except Exception as e:
                    raise RuntimeError(f"Failed to parse DOCX: {e}")
                    
            else:
                raise RuntimeError(f"Unsupported file type: {ext}")
                
        except RuntimeError:
            raise
        except Exception as e:
            raise RuntimeError(f"Unexpected error loading '{filename}': {e}")

    @lru_cache(maxsize=1000)
    def _clean_html_string(self, text: str) -> str:
        """Cached HTML cleaning for repeated strings."""
        import re
        # Remove HTML tags
        text = re.sub(r'<[^>]+>', '', text)
        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text)
        return text.strip()
    
    def remove_html(self, obj):
        """
        Recursively removes HTML tags from all string values with caching optimization.
        Args:
            obj: The JSON object (dict, list, or str).
        Returns:
            The cleaned object with HTML tags removed from strings.
        """
        if isinstance(obj, dict):
            return {k: self.remove_html(v) for k, v in obj.items()}
        elif isinstance(obj, list):
            return [self.remove_html(item) for item in obj]
        elif isinstance(obj, str):
            return self._clean_html_string(obj)
        else:
            return obj

    def text_to_chunks(self, text, max_chunk_size=500, overlap=50):
        """
        Splits plain text into chunks with optional overlap for better context preservation.
        Args:
            text (str): The text to chunk.
            max_chunk_size (int): Maximum size of each chunk.
            overlap (int): Number of characters to overlap between chunks.
        Returns:
            list: List of text chunks.
        """
        if len(text) <= max_chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + max_chunk_size
            
            # If not the last chunk, try to break at word boundary
            if end < len(text):
                # Find last space within chunk to avoid breaking words
                last_space = text.rfind(' ', start, end)
                if last_space > start:
                    end = last_space
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # Move start position with overlap
            start = end - overlap if end < len(text) else len(text)
            
        return chunks

    def json_to_chunks(self, obj, prefix="", max_chunk_size=500, parent=None, parent_map=None):
        """
        Recursively splits JSON data into text chunks for embedding with memory-efficient processing.
        Args:
            obj: The JSON object to chunk.
            prefix (str): Key path prefix for each chunk.
            max_chunk_size (int): Maximum size of each text chunk.
            parent: Parent object for mapping.
            parent_map (dict): Maps chunk text to its parent object.
        Returns:
            list: List of chunked text strings.
        """
        if parent_map is None:
            parent_map = {}
        chunks = []
        
        # Use generator for memory efficiency with large objects
        def _chunk_generator(obj, prefix, parent):
            if isinstance(obj, dict):
                for k, v in obj.items():
                    yield from _chunk_generator(v, prefix + k + ".", obj)
            elif isinstance(obj, list):
                # Process list items in batches for memory efficiency
                batch_size = 1000
                for i in range(0, len(obj), batch_size):
                    batch = obj[i:i + batch_size]
                    for idx, item in enumerate(batch):
                        actual_idx = i + idx
                        yield from _chunk_generator(item, prefix + str(actual_idx) + ".", obj)
            elif isinstance(obj, str):
                # Use optimized text chunking with overlap
                text_chunks = self.text_to_chunks(obj, max_chunk_size, overlap=50)
                for chunk in text_chunks:
                    chunk_text = f"{prefix[:-1]}: {chunk}"
                    yield chunk_text, parent
            else:
                chunk_text = f"{prefix[:-1]}: {obj}"
                yield chunk_text, parent
        
        # Process chunks in batches to manage memory
        for chunk_text, chunk_parent in _chunk_generator(obj, prefix, parent):
            chunks.append(chunk_text)
            parent_map[chunk_text] = chunk_parent
            
        return chunks

    def _setup_logging(self):
        """Setup logging for performance monitoring."""
        logging.basicConfig(level=logging.INFO)
        self.logger = logging.getLogger(__name__)

    @property
    def model(self):
        """Lazy loading of SentenceTransformer model."""
        if self._model is None:
            with self._model_lock:
                if self._model is None:
                    self.logger.info(f"Loading SentenceTransformer model: {self.model_name}")
                    self._model = SentenceTransformer(self.model_name)
        return self._model

    def _get_embedding_cache_key(self, text: str) -> str:
        """Generate cache key for embedding."""
        return hashlib.md5(f"{self.model_name}:{text}".encode()).hexdigest()

    def _get_cached_embedding(self, text: str) -> Optional[np.ndarray]:
        """Get cached embedding if available."""
        cache_key = self._get_embedding_cache_key(text)
        return self._embedding_cache.get(cache_key)

    def _cache_embedding(self, text: str, embedding: np.ndarray):
        """Cache embedding with LRU eviction."""
        if len(self._embedding_cache) >= self.cache_size:
            # Remove oldest item
            oldest_key = next(iter(self._embedding_cache))
            del self._embedding_cache[oldest_key]
        
        cache_key = self._get_embedding_cache_key(text)
        self._embedding_cache[cache_key] = embedding

    def encode_batch(self, texts: List[str]) -> np.ndarray:
        """Encode texts in batches with caching for better performance."""
        embeddings = []
        uncached_texts = []
        uncached_indices = []
        
        # Check cache first
        for i, text in enumerate(texts):
            cached_emb = self._get_cached_embedding(text)
            if cached_emb is not None:
                embeddings.append((i, cached_emb))
            else:
                uncached_texts.append(text)
                uncached_indices.append(i)
        
        # Encode uncached texts in batches
        if uncached_texts:
            self.logger.info(f"Encoding {len(uncached_texts)} texts in batches of {self.batch_size}")
            
            # Use tqdm progress bar for batch encoding
            with tqdm(total=len(uncached_texts), desc="Generating embeddings", unit="texts") as pbar:
                for i in range(0, len(uncached_texts), self.batch_size):
                    batch = uncached_texts[i:i + self.batch_size]
                    batch_embeddings = self.model.encode(batch, show_progress_bar=False)
                    
                    for j, emb in enumerate(batch_embeddings):
                        text_idx = uncached_indices[i + j]
                        text = uncached_texts[i + j]
                        embeddings.append((text_idx, emb))
                        self._cache_embedding(text, emb)
                    
                    pbar.update(len(batch))
        
        # Sort by original index and return embeddings
        embeddings.sort(key=lambda x: x[0])
        return np.array([emb for _, emb in embeddings])

    def search_vectors(self, filename: str, query: str, k: int = 10, score_threshold: float = None) -> List[Tuple[str, float, Any]]:
        """Search vectors with optimized retrieval and optional score filtering."""
        try:
            index, chunk_texts, parent_map, _ = self.load_vector_db(filename)
            
            # Encode query
            query_embedding = self.model.encode([query])
            
            # Search with more results initially for filtering
            search_k = min(k * 2, len(chunk_texts)) if score_threshold else k
            distances, indices = index.search(query_embedding.astype(np.float32), search_k)
            
            results = []
            for i, (dist, idx) in enumerate(zip(distances[0], indices[0])):
                if idx != -1:  # Valid result
                    similarity_score = 1 / (1 + dist)  # Convert L2 distance to similarity
                    
                    if score_threshold is None or similarity_score >= score_threshold:
                        chunk_text = chunk_texts[idx]
                        parent = parent_map.get(chunk_text, None)
                        results.append((chunk_text, similarity_score, parent))
                        
                        if len(results) >= k:
                            break
            
            return results
            
        except Exception as e:
            self.logger.error(f"Search failed for '{filename}': {e}")
            raise RuntimeError(f"Search failed for '{filename}': {e}")

    def delete_vectors(self, filename: str, chunk_indices: List[int]) -> bool:
        """Delete specific vectors from the database by rebuilding without them."""
        try:
            index, chunk_texts, parent_map, embeddings = self.load_vector_db(filename)
            
            # Create mask for chunks to keep
            keep_mask = np.ones(len(chunk_texts), dtype=bool)
            keep_mask[chunk_indices] = False
            
            # Filter chunks, embeddings, and parent map
            filtered_chunks = [chunk_texts[i] for i in range(len(chunk_texts)) if keep_mask[i]]
            filtered_embeddings = embeddings[keep_mask]
            filtered_parent_map = {chunk: parent for chunk, parent in parent_map.items() 
                                 if chunk in filtered_chunks}
            
            # Rebuild index with filtered data
            new_index = faiss.IndexFlatL2(filtered_embeddings.shape[1])
            new_index.add(filtered_embeddings.astype(np.float32))
            
            # Save updated database
            paths = self.get_vectordb_paths(filename)
            faiss.write_index(new_index, paths['index'])
            np.save(paths['npy'], np.array(filtered_chunks, dtype=object))
            np.save(paths['parentmap'], filtered_parent_map)
            np.save(paths['emb'], filtered_embeddings)
            
            self.logger.info(f"Deleted {len(chunk_indices)} vectors from '{filename}'")
            return True
            
        except Exception as e:
            self.logger.error(f"Vector deletion failed for '{filename}': {e}")
            raise RuntimeError(f"Vector deletion failed for '{filename}': {e}")

    def get_database_stats(self, filename: str) -> Dict[str, Any]:
        """Get statistics about the vector database."""
        try:
            paths = self.get_vectordb_paths(filename)
            if not os.path.exists(paths['index']):
                return {"exists": False}
            
            index, chunk_texts, parent_map, embeddings = self.load_vector_db(filename)
            
            stats = {
                "exists": True,
                "total_vectors": len(chunk_texts),
                "embedding_dimension": embeddings.shape[1] if len(embeddings) > 0 else 0,
                "index_size_mb": os.path.getsize(paths['index']) / (1024 * 1024),
                "chunks_size_mb": os.path.getsize(paths['npy']) / (1024 * 1024),
                "embeddings_size_mb": os.path.getsize(paths['emb']) / (1024 * 1024),
                "model_name": self.model_name
            }
            
            return stats
            
        except Exception as e:
            self.logger.error(f"Failed to get stats for '{filename}': {e}")
            return {"exists": False, "error": str(e)}

    def optimize_index(self, filename: str, use_ivf: bool = False, nlist: int = 100) -> bool:
        """Optimize the FAISS index for better search performance."""
        try:
            index, chunk_texts, parent_map, embeddings = self.load_vector_db(filename)
            
            if use_ivf and len(embeddings) > nlist * 39:  # IVF works better with more vectors
                # Use IVF index for better performance on large datasets
                quantizer = faiss.IndexFlatL2(embeddings.shape[1])
                new_index = faiss.IndexIVFFlat(quantizer, embeddings.shape[1], nlist)
                new_index.train(embeddings.astype(np.float32))
                new_index.add(embeddings.astype(np.float32))
                new_index.nprobe = min(10, nlist)  # Search 10 clusters by default
            else:
                # Use optimized flat index
                new_index = faiss.IndexFlatL2(embeddings.shape[1])
                new_index.add(embeddings.astype(np.float32))
            
            # Save optimized index
            paths = self.get_vectordb_paths(filename)
            faiss.write_index(new_index, paths['index'])
            
            self.logger.info(f"Optimized index for '{filename}' with {'IVF' if use_ivf else 'Flat'} index")
            return True
            
        except Exception as e:
            self.logger.error(f"Index optimization failed for '{filename}': {e}")
            raise RuntimeError(f"Index optimization failed for '{filename}': {e}")

    def vectors_exist(self, filename):
        """
        Check if vector database files already exist for a given filename.
        Args:
            filename (str): The file name to check.
        Returns:
            bool: True if all required vector files exist, False otherwise.
        """
        paths = self.get_vectordb_paths(filename)
        
        # Check if all required files exist and are not empty
        required_files = ['index', 'npy', 'parentmap']
        for key in required_files:
            path = paths[key]
            if not os.path.exists(path) or os.path.getsize(path) == 0:
                return False
        
        return True

    def get_vectordb_paths(self, filename):
        """
        Returns a dict of file paths for storing vector DB components for a given filename.
        Args:
            filename (str): The file name.
        Returns:
            dict: Paths for index, chunks, parent map, and embeddings.
        """
        base = os.path.join(self.vector_data_dir, f"{filename}_vectordb")
        return {
            'index': base + ".index",
            'npy': base + ".npy",
            'parentmap': base + "_parentmap.npy",
            'emb': base + ".emb.npy"
        }

    def build_vector_db(self, filename, data=None):
        """
        Builds a new vector database with optimized processing and error handling.
        Args:
            filename (str): The file name.
            data (optional): Pre-loaded data.
        Returns:
            bool: True if successful.
        Raises:
            RuntimeError: If building fails.
        """
        paths = self.get_vectordb_paths(filename)
        
        try:
            # Load data if not provided
            if data is None:
                self.logger.info(f"Loading file: {filename}")
                data = self.load_file(filename)
            
            # Process data based on file type
            ext = os.path.splitext(filename)[1].lower()
            parent_map = {}
            
            if ext == '.docx' and isinstance(data, dict) and 'chunks' in data:
                chunks = data['chunks']
                parent_map = {chunk: chunk for chunk in chunks}
            else:
                self.logger.info(f"Cleaning and chunking data for: {filename}")
                cleaned_data = self.remove_html(data)
                chunks = self.json_to_chunks(cleaned_data, parent_map=parent_map)
            
            if not chunks:
                raise RuntimeError(f"No chunks generated from file: {filename}")
            
            self.logger.info(f"Generated {len(chunks)} chunks from {filename}")
            
            # Generate embeddings with progress tracking
            self.logger.info(f"Generating embeddings for {len(chunks)} chunks")
            embeddings = self.encode_batch(chunks)
            
            if len(embeddings) != len(chunks):
                raise RuntimeError(f"Embedding count ({len(embeddings)}) doesn't match chunk count ({len(chunks)})")
            
            # Create and populate FAISS index
            self.logger.info(f"Building FAISS index with {embeddings.shape[1]} dimensions")
            index = faiss.IndexFlatL2(embeddings.shape[1])
            index.add(embeddings.astype(np.float32))
            
            # Save all components with error checking
            self.logger.info(f"Saving vector database components")
            
            try:
                faiss.write_index(index, paths['index'])
            except Exception as e:
                raise RuntimeError(f"Failed to save FAISS index: {e}")
            
            try:
                np.save(paths['npy'], np.array(chunks, dtype=object))
            except Exception as e:
                raise RuntimeError(f"Failed to save chunks: {e}")
            
            try:
                np.save(paths['parentmap'], parent_map)
            except Exception as e:
                raise RuntimeError(f"Failed to save parent map: {e}")
            
            try:
                np.save(paths['emb'], embeddings)
            except Exception as e:
                raise RuntimeError(f"Failed to save embeddings: {e}")
            
            # Verify saved files
            for key, path in paths.items():
                if not os.path.exists(path) or os.path.getsize(path) == 0:
                    raise RuntimeError(f"Failed to create valid {key} file: {path}")
            
            self.logger.info(f"Successfully built vector DB for '{filename}' with {len(chunks)} chunks")
            return True
            
        except RuntimeError:
            raise
        except Exception as e:
            raise RuntimeError(f"Unexpected error building vector DB for '{filename}': {e}")

    def load_vector_db(self, filename):
        """
        Loads an existing vector database with optimized error handling and validation.
        Args:
            filename (str): The file name.
        Returns:
            tuple: (FAISS index, chunk texts, parent map, chunk embeddings)
        Raises:
            RuntimeError: If loading fails or files are missing.
        """
        paths = self.get_vectordb_paths(filename)
        
        try:
            # Check for required files with detailed error messages
            missing_files = []
            for key, path in paths.items():
                if key != 'emb' and not os.path.exists(path):
                    missing_files.append(key)
            
            if missing_files:
                raise FileNotFoundError(f"Vector DB files missing: {missing_files}. Please build the vector DB first.")
            
            # Validate file integrity
            for key, path in paths.items():
                if key != 'emb' and os.path.exists(path):
                    if os.path.getsize(path) == 0:
                        raise RuntimeError(f"Vector DB file '{key}' is empty: {path}")
            
            # Load FAISS index with error handling
            try:
                index = faiss.read_index(paths['index'])
            except Exception as e:
                raise RuntimeError(f"Failed to load FAISS index: {e}")
            
            # Load chunk texts
            try:
                chunk_texts = list(np.load(paths['npy'], allow_pickle=True))
            except Exception as e:
                raise RuntimeError(f"Failed to load chunk texts: {e}")
            
            # Load parent map
            try:
                parent_map_data = np.load(paths['parentmap'], allow_pickle=True)
                parent_map = parent_map_data.item() if hasattr(parent_map_data, 'item') else parent_map_data
            except Exception as e:
                raise RuntimeError(f"Failed to load parent map: {e}")
            
            # Load or generate embeddings
            if os.path.exists(paths['emb']):
                try:
                    chunk_embs = np.load(paths['emb'])
                    # Validate embedding dimensions
                    if len(chunk_embs) != len(chunk_texts):
                        self.logger.warning("Embedding count mismatch, regenerating embeddings")
                        chunk_embs = self.encode_batch(chunk_texts)
                        np.save(paths['emb'], chunk_embs)
                except Exception as e:
                    self.logger.warning(f"Failed to load embeddings, regenerating: {e}")
                    chunk_embs = self.encode_batch(chunk_texts)
                    np.save(paths['emb'], chunk_embs)
            else:
                self.logger.info("Embeddings file missing, generating new embeddings")
                chunk_embs = self.encode_batch(chunk_texts)
                np.save(paths['emb'], chunk_embs)
            
            # Final validation
            if len(chunk_texts) != len(chunk_embs):
                raise RuntimeError("Mismatch between chunk count and embedding count")
            
            if index.ntotal != len(chunk_texts):
                raise RuntimeError("Mismatch between index size and chunk count")
            
            return index, chunk_texts, parent_map, chunk_embs
            
        except (FileNotFoundError, RuntimeError):
            raise
        except Exception as e:
            raise RuntimeError(f"Unexpected error loading vector DB for '{filename}': {e}")

    def rebuild_vector_db(self, filename):
        """
        Deletes and rebuilds the vector database with proper cleanup and validation.
        Args:
            filename (str): The file name.
        Returns:
            bool: True if successful.
        """
        paths = self.get_vectordb_paths(filename)
        
        # Clear embedding cache for this file
        keys_to_remove = []
        for key in self._embedding_cache:
            if filename in key:
                keys_to_remove.append(key)
        for key in keys_to_remove:
            del self._embedding_cache[key]
        
        # Remove all related files with better error handling
        removed_files = []
        failed_removals = []
        
        for ext in ['index', 'npy', 'parentmap', 'emb']:
            try:
                if os.path.exists(paths[ext]):
                    os.remove(paths[ext])
                    removed_files.append(ext)
            except FileNotFoundError:
                pass
            except Exception as e:
                failed_removals.append(f"{ext}: {e}")
                self.logger.warning(f"Could not remove {paths[ext]}: {e}")
        
        if removed_files:
            self.logger.info(f"Removed old vector DB files: {removed_files}")
        
        if failed_removals:
            self.logger.warning(f"Failed to remove some files: {failed_removals}")
        
        # Rebuild the database
        try:
            result = self.build_vector_db(filename)
            if result:
                self.logger.info(f"Successfully rebuilt vector DB for '{filename}'")
            return result
        except Exception as e:
            self.logger.error(f"Failed to rebuild vector DB for '{filename}': {e}")
            raise

    def get_all_pdf_files(self) -> List[str]:
        """
        Get all PDF files in the folder path.
        
        Returns:
            List[str]: List of PDF filenames in the folder.
        """
        pdf_files = []
        try:
            for filename in os.listdir(self.folder_path):
                if filename.lower().endswith('.pdf') and not filename.startswith('.'):
                    pdf_files.append(filename)
            return sorted(pdf_files)
        except Exception as e:
            self.logger.error(f"Error listing PDF files: {e}")
            return []

    def build_vector_db_for_all_files(self, file_extensions=None) -> Dict[str, bool]:
        """
        Build vector databases for all supported files in the folder.
        
        Args:
            file_extensions (list, optional): List of file extensions to process. 
                                            Defaults to ['.pdf', '.json', '.docx'].
        
        Returns:
            Dict[str, bool]: Dictionary mapping filenames to build success status.
        """
        if file_extensions is None:
            file_extensions = ['.pdf', '.json', '.docx']
        
        results = {}
        processed_files = []
        
        try:
            # Get all files with supported extensions
            for filename in os.listdir(self.folder_path):
                if any(filename.lower().endswith(ext) for ext in file_extensions) and not filename.startswith('.'):
                    processed_files.append(filename)
            
            if not processed_files:
                self.logger.warning(f"No files found with extensions {file_extensions}")
                return results
            
            self.logger.info(f"Found {len(processed_files)} files to process: {processed_files}")
            
            # Process each file with progress bar
            with tqdm(total=len(processed_files), desc="Processing files", unit="files") as pbar:
                for i, filename in enumerate(processed_files):
                    try:
                        pbar.set_description(f"Processing {filename}")
                        
                        # Check if vector DB already exists
                        stats = self.get_database_stats(filename)
                        if stats.get('exists', False):
                            self.logger.info(f"Vector DB already exists for '{filename}' with {stats['total_vectors']} vectors")
                            results[filename] = True
                        else:
                            self.logger.info(f"Building new vector DB for '{filename}'")
                            success = self.build_vector_db(filename)
                            results[filename] = success
                            
                            if success:
                                final_stats = self.get_database_stats(filename)
                                self.logger.info(f"Successfully built vector DB for '{filename}' with {final_stats['total_vectors']} vectors")
                            
                    except Exception as e:
                        self.logger.error(f"Failed to process '{filename}': {e}")
                        results[filename] = False
                    
                    pbar.update(1)
            
            # Summary
            successful = sum(1 for success in results.values() if success)
            self.logger.info(f"Processed {len(processed_files)} files: {successful} successful, {len(processed_files) - successful} failed")
            
            return results
            
        except Exception as e:
            self.logger.error(f"Error processing files: {e}")
            return results

    def delete_all_vector_databases(self) -> Dict[str, bool]:
        """
        Delete all vector databases in the vector data directory.
        
        Returns:
            Dict[str, bool]: Dictionary mapping filenames to deletion success status.
        """
        results = {}
        
        try:
            if not os.path.exists(self.vector_data_dir):
                self.logger.info("Vector data directory does not exist")
                return results
            
            # Find all vector database files
            vector_files = {}
            for filename in os.listdir(self.vector_data_dir):
                if filename.endswith('_vectordb.index'):
                    # Extract original filename
                    original_filename = filename.replace('_vectordb.index', '')
                    vector_files[original_filename] = True
            
            if not vector_files:
                self.logger.info("No vector databases found to delete")
                return results
            
            self.logger.info(f"Found {len(vector_files)} vector databases to delete")
            
            # Delete each vector database with progress bar
            with tqdm(total=len(vector_files), desc="Deleting vector databases", unit="DBs") as pbar:
                for original_filename in vector_files.keys():
                    try:
                        pbar.set_description(f"Deleting {original_filename}")
                        
                        paths = self.get_vectordb_paths(original_filename)
                        deleted_files = []
                        failed_deletions = []
                        
                        # Delete all related files
                        for file_type, file_path in paths.items():
                            try:
                                if os.path.exists(file_path):
                                    os.remove(file_path)
                                    deleted_files.append(file_type)
                            except Exception as e:
                                failed_deletions.append(f"{file_type}: {e}")
                        
                        if deleted_files:
                            self.logger.info(f"Deleted vector DB for '{original_filename}': {deleted_files}")
                            results[original_filename] = True
                        else:
                            self.logger.warning(f"No files found to delete for '{original_filename}'")
                            results[original_filename] = False
                        
                        if failed_deletions:
                            self.logger.warning(f"Failed to delete some files for '{original_filename}': {failed_deletions}")
                            
                    except Exception as e:
                        self.logger.error(f"Error deleting vector DB for '{original_filename}': {e}")
                        results[original_filename] = False
                    
                    pbar.update(1)
            
            # Clear embedding cache
            self._embedding_cache.clear()
            self.logger.info("Cleared embedding cache")
            
            # Summary
            successful = sum(1 for success in results.values() if success)
            self.logger.info(f"Deletion complete: {successful} successful, {len(results) - successful} failed")
            
            return results
            
        except Exception as e:
            self.logger.error(f"Error during vector database deletion: {e}")
            return results

    def refresh_all_vector_databases(self, file_extensions=None) -> Dict[str, bool]:
        """
        Delete all existing vector databases and rebuild them from scratch.
        
        Args:
            file_extensions (list, optional): List of file extensions to process.
                                            Defaults to ['.pdf', '.json', '.docx'].
        
        Returns:
            Dict[str, bool]: Dictionary mapping filenames to refresh success status.
        """
        self.logger.info("Starting vector database refresh process")
        
        # First, delete all existing vector databases
        deletion_results = self.delete_all_vector_databases()
        if deletion_results:
            self.logger.info(f"Deleted {len(deletion_results)} existing vector databases")
        
        # Then, rebuild all vector databases
        build_results = self.build_vector_db_for_all_files(file_extensions)
        
        self.logger.info("Vector database refresh process completed")
        return build_results

    def search_all_databases(self, query: str, k: int = 5, score_threshold: float = None) -> Dict[str, List[Tuple[str, float, Any]]]:
        """
        Search across all available vector databases for a given query.
        
        Args:
            query (str): The search query.
            k (int): Number of results to return per database.
            score_threshold (float, optional): Minimum similarity score threshold.
        
        Returns:
            Dict[str, List[Tuple[str, float, Any]]]: Dictionary mapping filenames to search results.
        """
        all_results = {}
        
        try:
            # Find all available vector databases
            if not os.path.exists(self.vector_data_dir):
                self.logger.warning("Vector data directory does not exist")
                return all_results
            
            vector_dbs = []
            for filename in os.listdir(self.vector_data_dir):
                if filename.endswith('_vectordb.index'):
                    original_filename = filename.replace('_vectordb.index', '')
                    vector_dbs.append(original_filename)
            
            if not vector_dbs:
                self.logger.warning("No vector databases found")
                return all_results
            
            self.logger.info(f"Searching across {len(vector_dbs)} databases for query: '{query}'")
            
            # Search each database with progress bar
            with tqdm(total=len(vector_dbs), desc="Searching databases", unit="DBs") as pbar:
                for db_name in vector_dbs:
                    try:
                        pbar.set_description(f"Searching {db_name}")
                        
                        results = self.search_vectors(db_name, query, k, score_threshold)
                        if results:
                            all_results[db_name] = results
                            self.logger.info(f"Found {len(results)} results in '{db_name}'")
                        else:
                            self.logger.info(f"No results found in '{db_name}'")
                            
                    except Exception as e:
                        self.logger.error(f"Error searching database '{db_name}': {e}")
                        all_results[db_name] = []
                    
                    pbar.update(1)
            
            return all_results
            
        except Exception as e:
            self.logger.error(f"Error in search_all_databases: {e}")
            return all_results

def test_vector_db_with_claude():
    """
    Interactive test function to demonstrate the optimized VectorDBManager with Claude Sonnet integration.
    Tests PDF processing, vector search, and enhanced Q&A using Claude's knowledge.
    """
    import os
    import time
    from dotenv import load_dotenv
    import anthropic
    
    # Load environment variables
    load_dotenv()
    
    print("\n=== VectorDBManager with Claude Sonnet Test ===")
    print("This test will demonstrate the optimized vector database functionality with AI-enhanced responses.\n")
    
    # Get PDF file path from user
    pdf_path = input("Please enter the full path to a PDF document: ").strip().strip('"').strip("'")
    
    if not os.path.exists(pdf_path):
        print(f"Error: File not found at {pdf_path}")
        return
    
    if not pdf_path.lower().endswith('.pdf'):
        print("Error: Please provide a PDF file.")
        return
    
    # Initialize Claude client
    try:
        anthropic_api_key = os.getenv('ANTHROPIC_API_KEY')
        if not anthropic_api_key:
            print("Error: ANTHROPIC_API_KEY not found in .env file")
            return
        
        claude_client = anthropic.Anthropic(api_key=anthropic_api_key)
        print("✓ Claude Sonnet API initialized successfully")
    except Exception as e:
        print(f"Error initializing Claude API: {e}")
        return
    
    # Initialize VectorDBManager
    try:
        folder_path = os.path.dirname(pdf_path)
        filename = os.path.basename(pdf_path)
        
        print(f"\nInitializing VectorDBManager...")
        db_manager = VectorDBManager(
            folder_path=folder_path,
            batch_size=16,  # Smaller batch for demo
            cache_size=500
        )
        print("✓ VectorDBManager initialized with optimized settings")
        
    except Exception as e:
        print(f"Error initializing VectorDBManager: {e}")
        return
    
    # Process PDF and build vector database
    try:
        print(f"\nProcessing PDF: {filename}")
        print("This may take a few moments depending on file size...")
        
        start_time = time.time()
        
        # Check if vector DB already exists
        stats = db_manager.get_database_stats(filename)
        if stats.get('exists', False):
            print(f"\nExisting vector database found:")
            print(f"  - Total vectors: {stats['total_vectors']}")
            print(f"  - Embedding dimension: {stats['embedding_dimension']}")
            print(f"  - Index size: {stats['index_size_mb']:.2f} MB")
            
            rebuild = input("\nRebuild vector database? (y/n): ").lower().startswith('y')
            if rebuild:
                print("Rebuilding vector database...")
                db_manager.rebuild_vector_db(filename)
        else:
            print("Building new vector database...")
            db_manager.build_vector_db(filename)
        
        processing_time = time.time() - start_time
        
        # Get final stats
        final_stats = db_manager.get_database_stats(filename)
        print(f"\n✓ Vector database ready!")
        print(f"  - Processing time: {processing_time:.2f} seconds")
        print(f"  - Total vectors: {final_stats['total_vectors']}")
        print(f"  - Embedding dimension: {final_stats['embedding_dimension']}")
        print(f"  - Model: {final_stats['model_name']}")
        
    except Exception as e:
        print(f"Error processing PDF: {e}")
        return
    
    # Interactive Q&A session
    print("\n=== Interactive Q&A Session ===")
    print("Ask questions about the PDF document. Type 'quit' to exit.\n")
    
    def get_enhanced_response(question: str, context_chunks: list) -> str:
        """
        Generate enhanced response using Claude Sonnet with vector context.
        """
        try:
            # Combine context from top matching chunks
            context = "\n\n".join([chunk[0] for chunk in context_chunks[:3]])
            
            prompt = f"""Based on the following context from a document, please answer the user's question comprehensively. 
            Combine the information from the document context with your knowledge to provide a detailed, accurate response.
            
            Document Context:
            {context}
            
            User Question: {question}
            
            Please provide a comprehensive answer that:
            1. Directly addresses the question using information from the document
            2. Adds relevant context or explanations from your general knowledge
            3. Clearly indicates what information comes from the document vs. general knowledge
            4. Is well-structured and easy to understand"""
            
            message = claude_client.messages.create(
                model="claude-3-5-haiku-20241022",
                max_tokens=1000,
                messages=[{"role": "user", "content": prompt}]
            )
            
            return message.content[0].text
            
        except Exception as e:
            return f"Error generating enhanced response: {e}"
    
    question_count = 0
    while True:
        try:
            question = input("\n> Your question: ").strip()
            
            if question.lower() in ['quit', 'exit', 'q']:
                print("\nGoodbye!")
                break
            
            if not question:
                print("Please enter a question.")
                continue
            
            question_count += 1
            print(f"\n[Question {question_count}] Searching vector database...")
            
            # Search vector database
            search_start = time.time()
            results = db_manager.search_vectors(
                filename=filename,
                query=question,
                k=5,  # Get top 5 matches
                score_threshold=0.3  # Filter low-relevance results
            )
            search_time = time.time() - search_start
            
            if not results:
                print("No relevant information found in the document for your question.")
                
                # Ask Claude without document context
                claude_start = time.time()
                try:
                    message = claude_client.messages.create(
                        model="claude-3-5-haiku-20241022",
                        max_tokens=500,
                        messages=[{"role": "user", "content": f"Please answer this question using your general knowledge: {question}"}]
                    )
                    claude_time = time.time() - claude_start
                    
                    print(f"\n[Claude Response - General Knowledge]")
                    print(f"Response time: {claude_time:.2f}s")
                    print(f"\n{message.content[0].text}")
                    
                except Exception as e:
                    print(f"Error getting Claude response: {e}")
                
                continue
            
            # Display search results
            print(f"Found {len(results)} relevant chunks (search time: {search_time:.2f}s)")
            print("\n[Top Matching Content from Document:]")
            for i, (chunk, score, _) in enumerate(results[:2], 1):
                print(f"\n{i}. Relevance: {score:.3f}")
                print(f"{chunk[:200]}{'...' if len(chunk) > 200 else ''}")
            
            # Generate enhanced response with Claude
            print(f"\n[Generating Enhanced Response with Claude Sonnet...]")
            claude_start = time.time()
            enhanced_response = get_enhanced_response(question, results)
            claude_time = time.time() - claude_start
            
            print(f"\n[Enhanced AI Response]")
            print(f"Response time: {claude_time:.2f}s")
            print(f"Context sources: {len(results)} document chunks + Claude's knowledge")
            print(f"\n{enhanced_response}")
            
            # Offer to show more details
            if len(results) > 2:
                show_more = input(f"\nShow {len(results) - 2} additional relevant chunks? (y/n): ").lower().startswith('y')
                if show_more:
                    for i, (chunk, score, _) in enumerate(results[2:], 3):
                        print(f"\n{i}. Relevance: {score:.3f}")
                        print(f"{chunk[:300]}{'...' if len(chunk) > 300 else ''}")
            
        except KeyboardInterrupt:
            print("\n\nSession interrupted. Goodbye!")
            break
        except Exception as e:
            print(f"\nError processing question: {e}")
            continue
    
    # Final statistics
    print(f"\n=== Session Summary ===")
    print(f"Questions answered: {question_count}")
    final_stats = db_manager.get_database_stats(filename)
    print(f"Document vectors: {final_stats['total_vectors']}")
    print(f"Vector database size: {final_stats['index_size_mb'] + final_stats['chunks_size_mb'] + final_stats['embeddings_size_mb']:.2f} MB")
    print(f"Embedding model: {final_stats['model_name']}")
    print(f"AI model=claude-3-5-haiku-20241022")


def demo_vector_db_with_claude():
    """
    Non-interactive demo version for testing the optimized VectorDBManager.
    """
    import os
    import time
    from dotenv import load_dotenv
    import anthropic
    
    # Load environment variables
    load_dotenv()
    
    print("\n=== VectorDBManager with Claude Sonnet Demo ===")
    print("Automated demo of optimized vector database functionality with AI-enhanced responses.\n")
    
    # Use sample PDF documents folder
    folder_path = "/Users/paragjain/dev-works/My-LLM-WorkingArea/sample-pdf-documents"
    
    if not os.path.exists(folder_path):
        print(f"Error: Sample PDF folder not found at {folder_path}")
        return
    
    print(f"Using sample PDF folder: {folder_path}")
    
    # Initialize Claude client
    try:
        anthropic_api_key = os.getenv('ANTHROPIC_API_KEY')
        if not anthropic_api_key:
            print("Error: ANTHROPIC_API_KEY not found in .env file")
            return
        
        claude_client = anthropic.Anthropic(api_key=anthropic_api_key)
        print("✓ Claude Sonnet API initialized successfully")
    except Exception as e:
        print(f"Error initializing Claude API: {e}")
        return
    
    # Initialize VectorDBManager
    try:
        print(f"\nInitializing VectorDBManager...")
        db_manager = VectorDBManager(
            folder_path=folder_path,
            batch_size=16,
            cache_size=500
        )
        print("✓ VectorDBManager initialized with optimized settings")
        
    except Exception as e:
        print(f"Error initializing VectorDBManager: {e}")
        return
    
    # Process all PDF files and build vector databases
    try:
        print(f"\nProcessing all PDF files in folder...")
        
        start_time = time.time()
        
        # Get all PDF files
        pdf_files = db_manager.get_all_pdf_files()
        if not pdf_files:
            print("No PDF files found in the folder")
            return
        
        print(f"Found {len(pdf_files)} PDF files: {pdf_files}")
        
        # Build vector databases for all files
        build_results = db_manager.build_vector_db_for_all_files(['.pdf'])
        
        processing_time = time.time() - start_time
        
        # Show results
        successful_files = [f for f, success in build_results.items() if success]
        failed_files = [f for f, success in build_results.items() if not success]
        
        print(f"\n✓ Vector database processing complete!")
        print(f"  - Processing time: {processing_time:.2f} seconds")
        print(f"  - Successful: {len(successful_files)} files")
        print(f"  - Failed: {len(failed_files)} files")
        
        if successful_files:
            print(f"  - Successfully processed: {successful_files}")
        if failed_files:
            print(f"  - Failed to process: {failed_files}")
        
        if not successful_files:
            print("No vector databases available for demo")
            return
        
    except Exception as e:
        print(f"Error processing PDF: {e}")
        return
    
    # Demo questions
    demo_questions = [
        "What is health psychology?",
        "What are the main factors that influence health?",
        "How does stress affect physical health?",
        "What are some effective stress management techniques?"
    ]
    
    print("\n=== Demo Q&A Session ===")
    print(f"Testing with {len(demo_questions)} sample questions...\n")
    
    def get_enhanced_response(question: str, all_results: dict) -> str:
        try:
            # Combine context from all databases
            all_chunks = []
            for db_name, results in all_results.items():
                for chunk, score, _ in results[:2]:  # Take top 2 from each DB
                    all_chunks.append(f"[From {db_name}] {chunk}")
            
            context = "\n\n".join(all_chunks[:5])  # Limit to top 5 overall
            
            prompt = f"""Based on the following context from multiple health psychology documents, please answer the user's question comprehensively.
            
            Document Context:
            {context}
            
            User Question: {question}
            
            Please provide a comprehensive answer that combines information from the documents with relevant general knowledge."""
            
            message = claude_client.messages.create(
                model="claude-3-5-haiku-20241022",
                max_tokens=800,
                messages=[{"role": "user", "content": prompt}]
            )
            
            return message.content[0].text
            
        except Exception as e:
            return f"Error generating enhanced response: {e}"
    
    # Process questions with progress bar
    with tqdm(total=len(demo_questions), desc="Processing Q&A", unit="questions") as pbar:
        for i, question in enumerate(demo_questions, 1):
            try:
                pbar.set_description(f"Q{i}: {question[:30]}...")
                
                print(f"\n{'='*60}")
                print(f"[Question {i}] {question}")
                print(f"{'='*60}")
                
                # Search across all vector databases
                search_start = time.time()
                all_results = db_manager.search_all_databases(
                    query=question,
                    k=3,
                    score_threshold=0.3
                )
                search_time = time.time() - search_start
                
                total_results = sum(len(results) for results in all_results.values())
                print(f"\n[Search Results] Found {total_results} relevant chunks across {len(all_results)} databases (search time: {search_time:.2f}s)")
                
                if all_results:
                    print("\n[Top Document Context from Multiple Sources:]")
                    result_count = 1
                    for db_name, results in all_results.items():
                        if results:
                            print(f"\n--- From {db_name} ---")
                            for chunk, score, _ in results[:2]:  # Show top 2 from each
                                print(f"\n{result_count}. Relevance Score: {score:.3f}")
                                print(f"{chunk[:250]}{'...' if len(chunk) > 250 else ''}")
                                result_count += 1
                    
                    # Generate enhanced response
                    print(f"\n[Generating Enhanced Response...]")
                    claude_start = time.time()
                    enhanced_response = get_enhanced_response(question, all_results)
                    claude_time = time.time() - claude_start
                    
                    print(f"\n[AI Response] (Response time: {claude_time:.2f}s)")
                    print(f"{enhanced_response}")
                else:
                    print("No relevant context found in any documents.")
                    
            except Exception as e:
                print(f"Error processing question: {e}")
            
            pbar.update(1)
    
    # Final statistics
    print(f"\n{'='*60}")
    print(f"=== Demo Complete ===")
    print(f"{'='*60}")
    
    # Get combined stats from all databases
    total_vectors = 0
    total_size = 0
    processed_files = 0
    
    for filename in successful_files:
        stats = db_manager.get_database_stats(filename)
        if stats.get('exists', False):
            total_vectors += stats['total_vectors']
            total_size += stats['index_size_mb'] + stats['chunks_size_mb'] + stats['embeddings_size_mb']
            processed_files += 1
    
    print(f"Questions processed: {len(demo_questions)}")
    print(f"Databases searched: {processed_files}")
    print(f"Total document vectors: {total_vectors}")
    print(f"Total DB size: {total_size:.2f} MB")
    print(f"Embedding model: {db_manager.model_name}")
    print(f"AI model: Claude-3.5-Haiku")
    print("\nDemo completed successfully!")




def refresh_vector_databases():
    """
    Function to refresh all vector databases.
    """
    import time
    
    print("\n=== Refresh Vector Databases ===")
    folder_path = "/Users/paragjain/dev-works/My-LLM-WorkingArea/sample-pdf-documents"
    
    if not os.path.exists(folder_path):
        print(f"Error: Sample PDF folder not found at {folder_path}")
        return
    
    try:
        db_manager = VectorDBManager(folder_path=folder_path)
        print("✓ VectorDBManager initialized")
        
        print("\nRefreshing all vector databases...")
        start_time = time.time()
        
        results = db_manager.refresh_all_vector_databases(['.pdf'])
        
        processing_time = time.time() - start_time
        
        successful = sum(1 for success in results.values() if success)
        failed = len(results) - successful
        
        print(f"\n✓ Refresh complete!")
        print(f"  - Processing time: {processing_time:.2f} seconds")
        print(f"  - Successfully refreshed: {successful} databases")
        print(f"  - Failed: {failed} databases")
        
        if results:
            for filename, success in results.items():
                status = "✓" if success else "✗"
                print(f"  {status} {filename}")
        
    except Exception as e:
        print(f"Error during refresh: {e}")


def delete_vector_databases():
    """
    Function to delete all vector databases.
    """
    print("\n=== Delete Vector Databases ===")
    folder_path = "/Users/paragjain/dev-works/My-LLM-WorkingArea/sample-pdf-documents"
    
    if not os.path.exists(folder_path):
        print(f"Error: Sample PDF folder not found at {folder_path}")
        return
    
    # Confirm deletion
    confirm = input("Are you sure you want to delete ALL vector databases? (type 'yes' to confirm): ").strip().lower()
    if confirm != 'yes':
        print("Deletion cancelled.")
        return
    
    try:
        db_manager = VectorDBManager(folder_path=folder_path)
        print("✓ VectorDBManager initialized")
        
        print("\nDeleting all vector databases...")
        results = db_manager.delete_all_vector_databases()
        
        successful = sum(1 for success in results.values() if success)
        failed = len(results) - successful
        
        print(f"\n✓ Deletion complete!")
        print(f"  - Successfully deleted: {successful} databases")
        print(f"  - Failed: {failed} databases")
        
        if results:
            for filename, success in results.items():
                status = "✓" if success else "✗"
                print(f"  {status} {filename}")
        else:
            print("  - No vector databases found to delete")
        
    except Exception as e:
        print(f"Error during deletion: {e}")


if __name__ == "__main__":
    # Ask user which version to run
    print("\n=== Vector DB Manager Options ===")
    print("1. Interactive test (asks for PDF path and questions)")
    print("2. Automated demo (uses all sample PDFs)")
    print("3. Refresh vector databases (delete and rebuild all)")
    print("4. Delete vector databases (remove all vector data)")
    
    try:
        choice = input("\nSelect option (1-4): ").strip()
        if choice == "1":
            test_vector_db_with_claude()
        elif choice == "2":
            demo_vector_db_with_claude()
        elif choice == "3":
            refresh_vector_databases()
        elif choice == "4":
            delete_vector_databases()
        else:
            print("Invalid choice. Running automated demo...")
            demo_vector_db_with_claude()
    except (KeyboardInterrupt, EOFError):
        print("\nExiting...")
